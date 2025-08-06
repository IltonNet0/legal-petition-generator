import sys
from docx import Document
from docx.shared import Cm
from docxtpl import DocxTemplate
from models.petition_data import petition_data
from utils.paragraph_style import Set_paragraph_style


# def FUTURE():
#     print("Welcome to the Legal Petition Generator")
    
#     # Choose input method
#     input_method = input("Would you like to input data via (1) Interface or (2) Spreadsheet? ")
    
#     if input_method == '1':
#         data = start_interface()
#     elif input_method == '2':
#         file_path = input("Please enter the path to the spreadsheet: ")
#         data = process_spreadsheet(file_path)
#     else:
#         print("Invalid option. Exiting.")
#         sys.exit(1)

#     # Generate the legal petition
#     petition = generate_petition(data)
    
#     # Output the petition
#     print("Generated Petition:")
#     print(petition)

# if __name__ == "__main__":
    # FUTURE()

def Petition_information():
    print('forneça as informações abaixo para gerar a petição:')

    data = petition_data()
    data.court = input('Vara: ').upper()
    data.judge_name = input('Nome do juiz: ').upper()
    data.client_name = input('Nome do cliente: ').upper()

    return data

new_case = Petition_information()

doc = DocxTemplate('src/petition/Template - Ana Eloisa.docx')
context = {
    'new_case': new_case
}

doc.render(context)
doc.save(f"Teste {new_case.client_name} - petition.docx")


################## TESTING DOCXTPL ##################

# apresetation = f"EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) {new_case.judge_name} DO TRABALHO DA VARA DO TRABALHO DE {new_case.court}"

# # petition.add_paragraph(apresetation, level=1)

# header = petition.add_heading()
# header.text = f"{apresetation}\n"
# # header.alignment = 1
# # header.text_size = 20

# paragraph = petition.add_paragraph()
# paragraph.text = 'Testando as configurações de parágrafo.'
# Set_paragraph_style(paragraph, 'justify', 12, 6, 11)

# petition.save('petition_test.docx')


print(f'\n{'Petição salva!'}\n')
